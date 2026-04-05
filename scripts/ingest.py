"""Ingest Word documents into Qdrant for RAG.

Usage:
    uv run python scripts/ingest.py <docs_directory>
    uv run python scripts/ingest.py data/documents/

Reads all .docx files from the given directory, splits them into chunks,
generates embeddings, and stores them in the Qdrant ``hr_documents`` collection.
Each file is tracked as a ``DocumentRecord`` in SQLite with chunk count and
indexing timestamp.

Only .docx files are processed (NFR-3: PDF scans are not used).
"""

from __future__ import annotations

import asyncio
import logging
import sys
import uuid
from datetime import UTC, datetime
from pathlib import Path

from docx import Document as DocxDocument
from langchain_core.documents import Document as LCDocument
from langchain_qdrant import QdrantVectorStore
from langchain_text_splitters import RecursiveCharacterTextSplitter
from qdrant_client import QdrantClient

# Allow running from project root
sys.path.insert(0, ".")

from app.config import Settings
from app.rag.indexer import prepare_chunks
from app.rag.retriever import build_embeddings
from app.storage.database import init_db
from app.storage.document_repo import DocumentRepository
from app.storage.models import DocumentRecord, DocumentStatus

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s  %(levelname)-8s  %(name)s  %(message)s",
)
logger = logging.getLogger(__name__)

CHUNK_SIZE = 1000
CHUNK_OVERLAP = 200


# ── Docx parsing ──────────────────────────────────────────────────


def _extract_sections(path: Path) -> list[tuple[str, str]]:
    """Extract ``(heading, body_text)`` pairs from a .docx file.

    Paragraphs with a ``Heading *`` style start a new section.  All
    subsequent paragraphs belong to the same section until the next
    heading.
    """
    doc = DocxDocument(str(path))
    sections: list[tuple[str, str]] = []
    current_heading = ""
    paragraphs: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        if para.style and para.style.name and para.style.name.startswith("Heading"):
            # Flush accumulated paragraphs as one section
            if paragraphs:
                sections.append((current_heading, "\n".join(paragraphs)))
                paragraphs = []
            current_heading = text

        paragraphs.append(text)

    # Flush remaining paragraphs
    if paragraphs:
        sections.append((current_heading, "\n".join(paragraphs)))

    return sections


def load_docx(path: Path) -> list[LCDocument]:
    """Parse a .docx file into chunked LangChain ``Document`` objects.

    Each chunk carries metadata: ``source`` (filename) and ``section``
    (nearest heading above the chunk).
    """
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        separators=["\n\n", "\n", ". ", " "],
    )

    sections = _extract_sections(path)
    documents: list[LCDocument] = []

    for heading, body in sections:
        chunks = splitter.split_text(body)
        for chunk in chunks:
            documents.append(
                LCDocument(
                    page_content=chunk,
                    metadata={
                        "source": path.name,
                        "section": heading,
                    },
                )
            )

    return documents


DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


# ── Main ingestion logic ──────────────────────────────────────────


async def ingest(docs_dir: Path, settings: Settings) -> int:
    """Ingest all .docx files from *docs_dir* into Qdrant.

    Creates a ``DocumentRecord`` in SQLite for every file and enriches
    each chunk with ``document_id``, ``chunk_id``, ``filename``, ``s3_key``,
    and ``is_search_enabled`` metadata.

    Returns the total number of stored chunks.
    """
    docx_files = sorted(docs_dir.glob("*.docx"))
    if not docx_files:
        logger.warning("No .docx files found in %s", docs_dir)
        return 0

    logger.info("Found %d .docx file(s) to ingest", len(docx_files))

    # Initialise SQLite
    await init_db(settings.db_path)
    repo = DocumentRepository(settings.db_path)

    # Load, chunk, and register each document
    all_docs: list[LCDocument] = []
    file_chunks: dict[str, list[LCDocument]] = {}

    for path in docx_files:
        logger.info("Processing %s ...", path.name)
        raw_chunks = load_docx(path)
        logger.info("  -> %d chunk(s)", len(raw_chunks))

        doc_id = uuid.uuid4().hex
        now = datetime.now(UTC)
        record = DocumentRecord(
            document_id=doc_id,
            filename=path.name,
            title=path.stem,
            s3_key=f"documents/{path.name}",
            mime_type=DOCX_MIME,
            size_bytes=path.stat().st_size,
            status=DocumentStatus.processing,
            created_at=now,
            updated_at=now,
        )
        await repo.create(record)

        enriched = prepare_chunks(
            raw_chunks,
            document_id=doc_id,
            filename=path.name,
            s3_key=f"documents/{path.name}",
        )
        all_docs.extend(enriched)
        file_chunks[doc_id] = enriched

    if not all_docs:
        logger.warning("No text extracted from any document")
        return 0

    logger.info("Total: %d chunk(s) from %d file(s)", len(all_docs), len(docx_files))

    # Build embeddings
    embeddings = build_embeddings(settings)

    # Recreate collection for a clean ingest
    client = QdrantClient(
        url=settings.qdrant_url,
        api_key=settings.qdrant_api_key,
    )
    try:
        collection = settings.qdrant_collection
        if client.collection_exists(collection):
            client.delete_collection(collection)
            logger.info("Deleted existing collection '%s'", collection)

        QdrantVectorStore.from_documents(
            documents=all_docs,
            embedding=embeddings,
            url=settings.qdrant_url,
            api_key=settings.qdrant_api_key,
            collection_name=collection,
        )
        logger.info(
            "Ingestion complete: %d chunk(s) stored in '%s'",
            len(all_docs),
            collection,
        )

        # Update document records with results
        now = datetime.now(UTC)
        for doc_id, chunks in file_chunks.items():
            await repo.update(
                doc_id,
                status=DocumentStatus.completed,
                chunk_count=len(chunks),
                indexed_at=now,
            )
    except Exception:
        for doc_id in file_chunks:
            await repo.update(
                doc_id,
                status=DocumentStatus.failed,
                error="Bulk ingestion failed",
            )
        raise
    finally:
        client.close()

    return len(all_docs)


# ── CLI ───────────────────────────────────────────────────────────


def main() -> None:
    if len(sys.argv) < 2:
        print(f"Usage: {sys.argv[0]} <docs_directory>", file=sys.stderr)
        sys.exit(1)

    docs_dir = Path(sys.argv[1])
    if not docs_dir.is_dir():
        print(f"Error: '{docs_dir}' is not a directory", file=sys.stderr)
        sys.exit(1)

    settings = Settings()
    total = asyncio.run(ingest(docs_dir, settings))
    if total:
        logger.info("Done — %d chunk(s) indexed.", total)
    else:
        logger.warning("Nothing was indexed.")


if __name__ == "__main__":
    main()
