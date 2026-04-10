"""Document parsing and chunking for the RAG pipeline.

Extracts text from .docx, .doc, and .xlsx files, splits into sections, and chunks
for embedding.  Used by both ``scripts/ingest.py`` and the admin upload flow.
"""

from __future__ import annotations

from pathlib import Path
from typing import Literal

import sharepoint2text
from docx import Document as DocxDocument
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph
from langchain_core.documents import Document as LCDocument
from langchain_core.embeddings import Embeddings
from langchain_experimental.text_splitter import SemanticChunker
from langchain_text_splitters import RecursiveCharacterTextSplitter

CHUNK_SIZE = 500
CHUNK_OVERLAP = 50

DOCX_MIME = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)


def _parse_heading_level(style_name: str | None) -> int:
    """Parse heading level from style name (e.g., 'Heading 2' -> 2).

    Returns 0 for non-heading styles or if parsing fails.
    """
    if not style_name:
        return 0
    if not style_name.startswith("Heading"):
        return 0
    # Extract number from "Heading N" or "HeadingN"
    parts = style_name.split()
    if len(parts) == 2 and parts[1].isdigit():
        return int(parts[1])
    return 0


def _extract_sections(path: Path) -> list[tuple[str, str, int, str]]:
    """Extract ``(heading, body_text, level, section_path)`` tuples from a .docx file.

    Paragraphs with a ``Heading *`` style start a new section.  All
    subsequent paragraphs belong to the same section until the next
    heading.

    Returns:
        List of tuples containing:
        - heading: The heading text (empty string for text before first heading)
        - body_text: The body text of the section
        - level: The heading level (0 for text before first heading)
        - section_path: Breadcrumb path like "Heading1 > Heading2 > Heading3"
    """
    doc = DocxDocument(str(path))
    sections: list[tuple[str, str, int, str]] = []
    current_heading = ""
    current_level = 0
    paragraphs: list[str] = []
    heading_stack: list[tuple[int, str]] = []  # (level, heading_text)

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style_name = para.style.name if para.style else None
        if style_name and style_name.startswith("Heading"):
            # Save previous section
            if paragraphs:
                section_path = " > ".join([h for _, h in heading_stack]) if heading_stack else ""
                sections.append(
                    (current_heading, "\n".join(paragraphs), current_level, section_path)
                )
                paragraphs = []

            # Parse heading level and update stack
            level = _parse_heading_level(style_name)
            current_heading = text
            current_level = level

            # Pop headings of same or higher level, then push new one
            while heading_stack and heading_stack[-1][0] >= level:
                heading_stack.pop()
            heading_stack.append((level, text))

        paragraphs.append(text)

    # Save final section
    if paragraphs:
        section_path = " > ".join([h for _, h in heading_stack]) if heading_stack else ""
        sections.append((current_heading, "\n".join(paragraphs), current_level, section_path))

    return sections


def load_docx(
    path: Path,
    *,
    chunk_size: int = CHUNK_SIZE,
    chunk_overlap: int = CHUNK_OVERLAP,
    strategy: str = "recursive",
    embeddings: Embeddings | None = None,
    breakpoint_threshold_type: Literal[
        "percentile", "standard_deviation", "interquartile", "gradient"
    ] = "percentile",
    breakpoint_threshold_amount: float | int = 95,
) -> list[LCDocument]:
    """Parse a .docx file into chunked LangChain ``Document`` objects.

    Each chunk carries metadata: ``source`` (filename) and ``section``
    (nearest heading above the chunk).

    Args:
        path: Path to the .docx file.
        chunk_size: Maximum chunk size for recursive strategy.
        chunk_overlap: Overlap between chunks for recursive strategy.
        strategy: Chunking strategy - "recursive" (default) or "semantic".
        embeddings: Embeddings model required for semantic chunking.
        breakpoint_threshold_type: Threshold type for semantic chunking.
        breakpoint_threshold_amount: Threshold amount for semantic chunking.

    Returns:
        List of LangChain Document objects.

    Raises:
        ValueError: If strategy is "semantic" and embeddings is None.
    """
    doc = DocxDocument(str(path))

    # Process document elements in order (paragraphs and tables interleaved)
    text_sections: list[tuple[str, str, int, str]] = []  # (heading, body, level, section_path)
    table_chunks: list[LCDocument] = []

    current_heading = ""
    current_level = 0
    paragraphs: list[str] = []
    heading_stack: list[tuple[int, str]] = []

    def _get_section_path() -> str:
        return " > ".join([h for _, h in heading_stack]) if heading_stack else ""

    def _save_current_section() -> None:
        nonlocal paragraphs
        if paragraphs:
            section_path = _get_section_path()
            text_sections.append(
                (current_heading, "\n".join(paragraphs), current_level, section_path)
            )
            paragraphs = []

    def _format_table_as_markdown(table: Table) -> str:
        """Format a docx table as a Markdown table."""
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)

        if not rows:
            return ""

        # Build markdown table
        lines = []
        # Header row
        header = " | ".join(rows[0])
        lines.append(f"| {header} |")
        # Separator
        lines.append("| " + " | ".join(["---"] * len(rows[0])) + " |")
        # Data rows
        for row in rows[1:]:
            # Pad row to match header length
            padded = row + [""] * (len(rows[0]) - len(row))
            lines.append(f"| {' | '.join(padded)} |")

        return "\n".join(lines)

    # Iterate through body elements in document order
    for element in doc.element.body:
        if element.tag == qn("w:p"):
            # Paragraph
            para = Paragraph(element, doc)
            text = para.text.strip()
            if not text:
                continue

            style_name = para.style.name if para.style else None
            if style_name and style_name.startswith("Heading"):
                _save_current_section()
                level = _parse_heading_level(style_name)
                current_heading = text
                current_level = level
                # Update heading stack
                while heading_stack and heading_stack[-1][0] >= level:
                    heading_stack.pop()
                heading_stack.append((level, text))
            else:
                paragraphs.append(text)

        elif element.tag == qn("w:tbl"):
            # Table - save current text section first
            _save_current_section()

            table = Table(element, doc)
            md_table = _format_table_as_markdown(table)
            if md_table:
                section_path = _get_section_path()
                table_chunks.append(
                    LCDocument(
                        page_content=md_table,
                        metadata={
                            "source": path.name,
                            "section": current_heading,
                            "section_level": current_level,
                            "section_path": section_path,
                            "is_table": True,
                        },
                    )
                )

    # Save final text section
    _save_current_section()

    # Process text sections with chunking strategy
    if strategy == "recursive":
        splitter = RecursiveCharacterTextSplitter.from_tiktoken_encoder(
            encoding_name="cl100k_base",
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=["\n\n", "\n", ". ", " "],
        )

        documents: list[LCDocument] = []
        for heading, body, level, section_path in text_sections:
            chunks = splitter.split_text(body)
            for chunk in chunks:
                documents.append(
                    LCDocument(
                        page_content=chunk,
                        metadata={
                            "source": path.name,
                            "section": heading,
                            "section_level": level,
                            "section_path": section_path,
                        },
                    )
                )
        # Add table chunks (not split further)
        documents.extend(table_chunks)
        return documents

    if strategy == "semantic":
        if embeddings is None:
            raise ValueError("embeddings are required for semantic chunking strategy")

        # Build full text with section offsets tracking
        section_offsets: list[tuple[int, int, str, int, str]] = []
        full_text_parts: list[str] = []
        current_offset = 0

        for heading, body, level, section_path in text_sections:
            section_text = body
            start_offset = current_offset
            end_offset = current_offset + len(section_text)
            section_offsets.append((start_offset, end_offset, heading, level, section_path))
            full_text_parts.append(section_text)
            current_offset = end_offset + 2  # +2 for "\n\n" separator

        full_text = "\n\n".join(full_text_parts)

        chunker = SemanticChunker(
            embeddings,
            breakpoint_threshold_type=breakpoint_threshold_type,
            breakpoint_threshold_amount=breakpoint_threshold_amount,
        )

        semantic_chunks: list[LCDocument] = chunker.create_documents([full_text])

        semantic_documents: list[LCDocument] = []
        for sem_chunk in semantic_chunks:
            chunk_text = sem_chunk.page_content
            # Find chunk position in full text
            chunk_start = full_text.find(chunk_text)
            if chunk_start == -1:
                chunk_start = 0
            chunk_end = chunk_start + len(chunk_text)

            # Find section with largest overlap
            best_heading = ""
            best_level = 0
            best_section_path = ""
            best_overlap = 0
            for start, end, heading, level, section_path in section_offsets:
                overlap_start = max(start, chunk_start)
                overlap_end = min(end, chunk_end)
                overlap = max(0, overlap_end - overlap_start)
                if overlap > best_overlap:
                    best_overlap = overlap
                    best_heading = heading
                    best_level = level
                    best_section_path = section_path

            semantic_documents.append(
                LCDocument(
                    page_content=chunk_text,
                    metadata={
                        "source": path.name,
                        "section": best_heading,
                        "section_level": best_level,
                        "section_path": best_section_path,
                    },
                )
            )

        # Add table chunks (not split further)
        semantic_documents.extend(table_chunks)
        return semantic_documents

    raise ValueError(f"Unknown chunking strategy: {strategy}")


def load_doc(
    path: Path,
    *,
    chunk_size: int = CHUNK_SIZE,
    chunk_overlap: int = CHUNK_OVERLAP,
    strategy: str = "recursive",
    embeddings: Embeddings | None = None,
    breakpoint_threshold_type: Literal[
        "percentile", "standard_deviation", "interquartile", "gradient"
    ] = "percentile",
    breakpoint_threshold_amount: float | int = 95,
) -> list[LCDocument]:
    """Parse a legacy .doc file into chunked LangChain ``Document`` objects.

    Since .doc files lack structured heading styles, the entire text is treated
    as one section with the filename stem as the heading for recursive strategy,
    or empty section for semantic strategy.

    Each chunk carries metadata: ``source`` (filename) and ``section``
    (filename stem for recursive, empty for semantic).

    Args:
        path: Path to the .doc file.
        chunk_size: Maximum chunk size for recursive strategy.
        chunk_overlap: Overlap between chunks for recursive strategy.
        strategy: Chunking strategy - "recursive" (default) or "semantic".
        embeddings: Embeddings model required for semantic chunking.
        breakpoint_threshold_type: Threshold type for semantic chunking.
        breakpoint_threshold_amount: Threshold amount for semantic chunking.

    Returns:
        List of LangChain Document objects.

    Raises:
        ValueError: If strategy is "semantic" and embeddings is None.
    """
    # Use sharepoint2text to extract text from legacy .doc files
    # timeout_seconds=0 disables signal-based timeout (requires main thread)
    # since this function may be called via asyncio.to_thread() in a worker thread
    result = next(sharepoint2text.read_file(str(path), timeout_seconds=0))
    text = result.get_full_text()

    if strategy == "recursive":
        splitter = RecursiveCharacterTextSplitter.from_tiktoken_encoder(
            encoding_name="cl100k_base",
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=["\n\n", "\n", ". ", " "],
        )

        section_heading = path.stem
        chunks = splitter.split_text(text)
        doc_chunks: list[LCDocument] = []

        for chunk in chunks:
            doc_chunks.append(
                LCDocument(
                    page_content=chunk,
                    metadata={
                        "source": path.name,
                        "section": section_heading,
                    },
                )
            )

        return doc_chunks

    if strategy == "semantic":
        if embeddings is None:
            raise ValueError("embeddings are required for semantic chunking strategy")

        chunker = SemanticChunker(
            embeddings,
            breakpoint_threshold_type=breakpoint_threshold_type,
            breakpoint_threshold_amount=breakpoint_threshold_amount,
        )

        semantic_chunks: list[LCDocument] = chunker.create_documents([text])

        documents: list[LCDocument] = []
        for sem_chunk in semantic_chunks:
            documents.append(
                LCDocument(
                    page_content=sem_chunk.page_content,
                    metadata={
                        "source": path.name,
                        "section": "",
                    },
                )
            )

        return documents

    raise ValueError(f"Unknown chunking strategy: {strategy}")


def load_xlsx(
    path: Path,
    *,
    chunk_size: int = CHUNK_SIZE,
    chunk_overlap: int = CHUNK_OVERLAP,
    strategy: str = "recursive",
    embeddings: Embeddings | None = None,
    breakpoint_threshold_type: Literal[
        "percentile", "standard_deviation", "interquartile", "gradient"
    ] = "percentile",
    breakpoint_threshold_amount: float | int = 95,
) -> list[LCDocument]:
    """Parse an .xlsx spreadsheet into chunked LangChain ``Document`` objects.

    Each worksheet becomes a section (sheet name as heading). Rows are formatted
    with `` | `` separators to preserve column structure. Empty rows are skipped.

    The first non-empty row of each sheet is extracted as column headers and
    prepended to each chunk for context.

    Each chunk carries metadata: ``source`` (filename), ``section`` (sheet name),
    and ``column_headers`` (the header row as a formatted string).

    Args:
        path: Path to the .xlsx file.
        chunk_size: Maximum chunk size for recursive strategy.
        chunk_overlap: Overlap between chunks for recursive strategy.
        strategy: Chunking strategy - "recursive" (default) or "semantic".
        embeddings: Embeddings model required for semantic chunking.
        breakpoint_threshold_type: Threshold type for semantic chunking.
        breakpoint_threshold_amount: Threshold amount for semantic chunking.

    Returns:
        List of LangChain Document objects.

    Raises:
        ValueError: If strategy is "semantic" and embeddings is None.
    """
    import openpyxl

    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    sections: list[tuple[str, str, str]] = []  # (sheet_name, body, column_headers)

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows_text: list[str] = []
        column_headers = ""
        header_found = False

        for row in ws.iter_rows(values_only=True):
            # Skip completely empty rows
            cells = [str(cell) if cell is not None else "" for cell in row]
            if not any(c.strip() for c in cells):
                continue

            row_text = " | ".join(cells)

            # First non-empty row becomes the header
            if not header_found:
                column_headers = row_text
                header_found = True

            rows_text.append(row_text)

        if rows_text:
            sections.append((sheet_name, "\n".join(rows_text), column_headers))

    wb.close()

    if strategy == "recursive":
        splitter = RecursiveCharacterTextSplitter.from_tiktoken_encoder(
            encoding_name="cl100k_base",
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=["\n\n", "\n", ". ", " "],
        )

        documents: list[LCDocument] = []
        for sheet_name, body, column_headers in sections:
            chunks = splitter.split_text(body)
            for chunk in chunks:
                # Prepend headers to chunk content for context
                chunk_with_headers = f"{column_headers}\n{chunk}"
                documents.append(
                    LCDocument(
                        page_content=chunk_with_headers,
                        metadata={
                            "source": path.name,
                            "section": sheet_name,
                            "column_headers": column_headers,
                        },
                    )
                )
        return documents

    if strategy == "semantic":
        if embeddings is None:
            raise ValueError("embeddings are required for semantic chunking strategy")

        # Build full text with section offsets tracking
        section_offsets: list[tuple[int, int, str, str]] = []
        full_text_parts: list[str] = []
        current_offset = 0

        for sheet_name, body, column_headers in sections:
            section_text = body
            start_offset = current_offset
            end_offset = current_offset + len(section_text)
            section_offsets.append((start_offset, end_offset, sheet_name, column_headers))
            full_text_parts.append(section_text)
            current_offset = end_offset + 2  # +2 for "\n\n" separator

        full_text = "\n\n".join(full_text_parts)

        chunker = SemanticChunker(
            embeddings,
            breakpoint_threshold_type=breakpoint_threshold_type,
            breakpoint_threshold_amount=breakpoint_threshold_amount,
        )

        semantic_chunks: list[LCDocument] = chunker.create_documents([full_text])

        semantic_documents: list[LCDocument] = []
        for sem_chunk in semantic_chunks:
            chunk_text = sem_chunk.page_content
            # Find chunk position in full text
            chunk_start = full_text.find(chunk_text)
            if chunk_start == -1:
                chunk_start = 0
            chunk_end = chunk_start + len(chunk_text)

            # Find section with largest overlap
            best_sheet_name = ""
            best_column_headers = ""
            best_overlap = 0
            for start, end, sheet_name, column_headers in section_offsets:
                overlap_start = max(start, chunk_start)
                overlap_end = min(end, chunk_end)
                overlap = max(0, overlap_end - overlap_start)
                if overlap > best_overlap:
                    best_overlap = overlap
                    best_sheet_name = sheet_name
                    best_column_headers = column_headers

            # Prepend headers to chunk content for context
            chunk_with_headers = f"{best_column_headers}\n{chunk_text}"
            semantic_documents.append(
                LCDocument(
                    page_content=chunk_with_headers,
                    metadata={
                        "source": path.name,
                        "section": best_sheet_name,
                        "column_headers": best_column_headers,
                    },
                )
            )

        return semantic_documents

    raise ValueError(f"Unknown chunking strategy: {strategy}")



def load_document(
    path: Path,
    *,
    chunk_size: int = CHUNK_SIZE,
    chunk_overlap: int = CHUNK_OVERLAP,
    strategy: str = "recursive",
    embeddings: Embeddings | None = None,
    breakpoint_threshold_type: Literal[
        "percentile", "standard_deviation", "interquartile", "gradient"
    ] = "percentile",
    breakpoint_threshold_amount: float | int = 95,
) -> list[LCDocument]:
    """Parse a document file into chunked LangChain ``Document`` objects.

    Dispatches to the appropriate loader based on file extension:
    - .docx -> load_docx
    - .doc -> load_doc
    - .xlsx -> load_xlsx

    Args:
        path: Path to the document file.
        chunk_size: Maximum chunk size for recursive strategy.
        chunk_overlap: Overlap between chunks for recursive strategy.
        strategy: Chunking strategy - "recursive" (default) or "semantic".
        embeddings: Embeddings model required for semantic chunking.
        breakpoint_threshold_type: Threshold type for semantic chunking.
        breakpoint_threshold_amount: Threshold amount for semantic chunking.

    Raises:
        ValueError: If the file extension is not supported.
    """
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return load_docx(
            path,
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            strategy=strategy,
            embeddings=embeddings,
            breakpoint_threshold_type=breakpoint_threshold_type,
            breakpoint_threshold_amount=breakpoint_threshold_amount,
        )
    elif suffix == ".doc":
        return load_doc(
            path,
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            strategy=strategy,
            embeddings=embeddings,
            breakpoint_threshold_type=breakpoint_threshold_type,
            breakpoint_threshold_amount=breakpoint_threshold_amount,
        )
    elif suffix == ".xlsx":
        return load_xlsx(
            path,
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            strategy=strategy,
            embeddings=embeddings,
            breakpoint_threshold_type=breakpoint_threshold_type,
            breakpoint_threshold_amount=breakpoint_threshold_amount,
        )
    else:
        raise ValueError(f"Unsupported file extension: {suffix}")
