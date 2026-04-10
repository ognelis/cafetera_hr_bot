"""Tests for app.rag.parser — document parsing functions."""

from __future__ import annotations

from pathlib import Path
from unittest.mock import MagicMock, patch

import pytest
from docx import Document as DocxDocument
from langchain_core.documents import Document as LCDocument

from app.rag.parser import load_doc, load_document, load_docx, load_xlsx

# ── load_doc ──────────────────────────────────────────────────────


class TestLoadDoc:
    def test_load_doc_returns_documents(self, tmp_path: Path):
        doc_path = tmp_path / "test.doc"

        with patch("app.rag.parser.sharepoint2text.read_file") as mock_read_file:
            mock_result = MagicMock()
            mock_result.get_full_text.return_value = (
                "First paragraph.\n\nSecond paragraph with more content."
            )
            mock_read_file.return_value = iter([mock_result])
            result = load_doc(doc_path)

        assert len(result) > 0
        assert all(isinstance(doc, LCDocument) for doc in result)
        mock_read_file.assert_called_once_with(str(doc_path), timeout_seconds=0)

    def test_load_doc_metadata(self, tmp_path: Path):
        doc_path = tmp_path / "my_document.doc"

        with patch("app.rag.parser.sharepoint2text.read_file") as mock_read_file:
            mock_result = MagicMock()
            mock_result.get_full_text.return_value = "Some content here."
            mock_read_file.return_value = iter([mock_result])
            result = load_doc(doc_path)

        assert len(result) > 0
        for doc in result:
            assert doc.metadata["source"] == "my_document.doc"
            assert doc.metadata["section"] == "my_document"

    def test_load_doc_multi_paragraph_content(self, tmp_path: Path):
        doc_path = tmp_path / "test.doc"
        sample_text = "Paragraph one.\n\nParagraph two.\n\nParagraph three with more text."

        with patch("app.rag.parser.sharepoint2text.read_file") as mock_read_file:
            mock_result = MagicMock()
            mock_result.get_full_text.return_value = sample_text
            mock_read_file.return_value = iter([mock_result])
            result = load_doc(doc_path)

        assert len(result) > 0
        combined_content = " ".join(doc.page_content for doc in result)
        assert "Paragraph one" in combined_content
        assert "Paragraph two" in combined_content
        assert "Paragraph three" in combined_content


# ── load_document dispatcher ──────────────────────────────────────


class TestLoadDocumentDispatcher:
    def test_load_document_docx(self, tmp_path: Path):
        docx_path = tmp_path / "test.docx"
        doc = DocxDocument()
        doc.add_heading("Test Heading", level=1)
        doc.add_paragraph("Test paragraph content.")
        doc.save(docx_path)

        result = load_document(docx_path)

        assert len(result) > 0
        assert all(isinstance(doc, LCDocument) for doc in result)
        for doc in result:
            assert doc.metadata["source"] == "test.docx"

    def test_load_document_doc(self, tmp_path: Path):
        doc_path = tmp_path / "test.doc"

        with patch("app.rag.parser.sharepoint2text.read_file") as mock_read_file:
            mock_result = MagicMock()
            mock_result.get_full_text.return_value = "Legacy document content."
            mock_read_file.return_value = iter([mock_result])
            result = load_document(doc_path)

        assert len(result) > 0
        assert all(isinstance(doc, LCDocument) for doc in result)
        mock_read_file.assert_called_once_with(str(doc_path), timeout_seconds=0)

    def test_load_document_unsupported_extension(self, tmp_path: Path):
        txt_path = tmp_path / "test.txt"
        txt_path.write_text("Some text content.")

        with pytest.raises(ValueError) as exc_info:
            load_document(txt_path)

        assert "Unsupported file extension: .txt" in str(exc_info.value)


# ── load_xlsx ─────────────────────────────────────────────────────


def _create_mock_workbook(sheet_data: dict[str, list[tuple]]) -> MagicMock:
    """Create a mock openpyxl workbook with specified sheet data.

    Args:
        sheet_data: Dict mapping sheet name to list of row tuples.
    """
    mock_wb = MagicMock()
    mock_wb.sheetnames = list(sheet_data.keys())

    def get_sheet(name: str):
        mock_ws = MagicMock()
        mock_ws.iter_rows.return_value = sheet_data[name]
        return mock_ws

    mock_wb.__getitem__ = lambda self, name: get_sheet(name)
    return mock_wb


class TestLoadXlsx:
    def test_load_xlsx_returns_documents(self, tmp_path: Path):
        xlsx_path = tmp_path / "test.xlsx"

        sheet_data = {
            "Sheet1": [
                ("A1", "B1", "C1"),
                ("A2", "B2", "C2"),
            ]
        }
        mock_wb = _create_mock_workbook(sheet_data)

        with patch("openpyxl.load_workbook") as mock_load:
            mock_load.return_value = mock_wb
            result = load_xlsx(xlsx_path)

        assert len(result) > 0
        assert all(isinstance(doc, LCDocument) for doc in result)
        mock_load.assert_called_once_with(xlsx_path, read_only=True, data_only=True)
        mock_wb.close.assert_called_once()

    def test_load_xlsx_metadata(self, tmp_path: Path):
        xlsx_path = tmp_path / "my_spreadsheet.xlsx"

        sheet_data = {
            "Sales": [
                ("Product", "Revenue"),
                ("Widget", "1000"),
            ]
        }
        mock_wb = _create_mock_workbook(sheet_data)

        with patch("openpyxl.load_workbook") as mock_load:
            mock_load.return_value = mock_wb
            result = load_xlsx(xlsx_path)

        assert len(result) > 0
        for doc in result:
            assert doc.metadata["source"] == "my_spreadsheet.xlsx"
            assert doc.metadata["section"] == "Sales"

    def test_load_xlsx_multi_row_content(self, tmp_path: Path):
        xlsx_path = tmp_path / "test.xlsx"

        sheet_data = {
            "Data": [
                ("Row1Col1", "Row1Col2"),
                ("Row2Col1", "Row2Col2"),
                ("Row3Col1", "Row3Col2"),
            ]
        }
        mock_wb = _create_mock_workbook(sheet_data)

        with patch("openpyxl.load_workbook") as mock_load:
            mock_load.return_value = mock_wb
            result = load_xlsx(xlsx_path)

        assert len(result) > 0
        combined_content = " ".join(doc.page_content for doc in result)
        assert "Row1Col1" in combined_content
        assert "Row2Col2" in combined_content
        assert "Row3Col1" in combined_content

    def test_load_xlsx_skips_empty_rows(self, tmp_path: Path):
        xlsx_path = tmp_path / "test.xlsx"

        sheet_data = {
            "Sheet1": [
                ("A1", "B1"),
                (None, None),  # Empty row
                ("A3", "B3"),
            ]
        }
        mock_wb = _create_mock_workbook(sheet_data)

        with patch("openpyxl.load_workbook") as mock_load:
            mock_load.return_value = mock_wb
            result = load_xlsx(xlsx_path)

        assert len(result) > 0
        combined_content = " ".join(doc.page_content for doc in result)
        assert "A1" in combined_content
        assert "A3" in combined_content

    def test_load_xlsx_multiple_sheets(self, tmp_path: Path):
        xlsx_path = tmp_path / "test.xlsx"

        sheet_data = {
            "Sheet1": [("Data1",), ("Data2",)],
            "Sheet2": [("Data3",), ("Data4",)],
        }
        mock_wb = _create_mock_workbook(sheet_data)

        with patch("openpyxl.load_workbook") as mock_load:
            mock_load.return_value = mock_wb
            result = load_xlsx(xlsx_path)

        assert len(result) > 0
        sheet1_docs = [doc for doc in result if doc.metadata["section"] == "Sheet1"]
        sheet2_docs = [doc for doc in result if doc.metadata["section"] == "Sheet2"]
        assert len(sheet1_docs) > 0
        assert len(sheet2_docs) > 0


# ── load_document dispatcher for xlsx ─────────────────────────────


class TestLoadDocumentXlsxDispatcher:
    def test_load_document_xlsx(self, tmp_path: Path):
        xlsx_path = tmp_path / "test.xlsx"

        sheet_data = {
            "Sheet1": [("Spreadsheet content from xlsx.",)],
        }
        mock_wb = _create_mock_workbook(sheet_data)

        with patch("openpyxl.load_workbook") as mock_load:
            mock_load.return_value = mock_wb
            result = load_document(xlsx_path)

        assert len(result) > 0
        assert all(isinstance(doc, LCDocument) for doc in result)
        mock_load.assert_called_once_with(xlsx_path, read_only=True, data_only=True)


# ── load_docx heading level and section path tests ─────────────────


class TestLoadDocxHeadingLevels:
    """Tests for heading level tracking and section path breadcrumbs."""

    def test_load_docx_heading_levels(self, tmp_path: Path):
        """Create a .docx with Heading 1, Heading 2, and body text.
        Verify chunks have correct section_level and section_path metadata."""
        docx_path = tmp_path / "heading_levels.docx"
        doc = DocxDocument()

        # Heading 1
        doc.add_heading("Main Section", level=1)
        doc.add_paragraph("This is content under the main section.")

        # Heading 2
        doc.add_heading("Subsection", level=2)
        doc.add_paragraph("This is content under the subsection.")

        doc.save(docx_path)

        result = load_docx(docx_path)

        # Find chunks for each section
        main_section_chunks = [
            d for d in result
            if d.metadata.get("section") == "Main Section"
        ]
        subsection_chunks = [
            d for d in result
            if d.metadata.get("section") == "Subsection"
        ]

        assert len(main_section_chunks) > 0
        assert len(subsection_chunks) > 0

        # Verify section_level
        for chunk in main_section_chunks:
            assert chunk.metadata["section_level"] == 1
            assert chunk.metadata["section_path"] == "Main Section"

        for chunk in subsection_chunks:
            assert chunk.metadata["section_level"] == 2
            assert chunk.metadata["section_path"] == "Main Section > Subsection"

    def test_load_docx_section_path_breadcrumbs(self, tmp_path: Path):
        """Create a .docx with nested headings (H1 > H2 > H3, then new H1).
        Verify section_path resets correctly when a higher-level heading appears."""
        docx_path = tmp_path / "breadcrumb_reset.docx"
        doc = DocxDocument()

        # First H1 > H2 > H3 hierarchy
        doc.add_heading("Chapter 1", level=1)
        doc.add_paragraph("Chapter 1 introduction text.")

        doc.add_heading("Section 1.1", level=2)
        doc.add_paragraph("Section 1.1 content.")

        doc.add_heading("Subsection 1.1.1", level=3)
        doc.add_paragraph("Subsection 1.1.1 content.")

        # New H1 - should reset the path
        doc.add_heading("Chapter 2", level=1)
        doc.add_paragraph("Chapter 2 introduction text.")

        doc.save(docx_path)

        result = load_docx(docx_path)

        # Find chunks for each section
        chapter1_chunks = [
            d for d in result
            if d.metadata.get("section") == "Chapter 1"
        ]
        section11_chunks = [
            d for d in result
            if d.metadata.get("section") == "Section 1.1"
        ]
        subsection111_chunks = [
            d for d in result
            if d.metadata.get("section") == "Subsection 1.1.1"
        ]
        chapter2_chunks = [
            d for d in result
            if d.metadata.get("section") == "Chapter 2"
        ]

        assert len(chapter1_chunks) > 0
        assert len(section11_chunks) > 0
        assert len(subsection111_chunks) > 0
        assert len(chapter2_chunks) > 0

        # Verify section paths
        for chunk in chapter1_chunks:
            assert chunk.metadata["section_path"] == "Chapter 1"
            assert chunk.metadata["section_level"] == 1

        for chunk in section11_chunks:
            assert chunk.metadata["section_path"] == "Chapter 1 > Section 1.1"
            assert chunk.metadata["section_level"] == 2

        for chunk in subsection111_chunks:
            assert chunk.metadata["section_path"] == "Chapter 1 > Section 1.1 > Subsection 1.1.1"
            assert chunk.metadata["section_level"] == 3

        # Chapter 2 should have reset path
        for chunk in chapter2_chunks:
            assert chunk.metadata["section_path"] == "Chapter 2"
            assert chunk.metadata["section_level"] == 1


# ── load_docx table tests ─────────────────────────────────────────


class TestLoadDocxTables:
    """Tests for table extraction as atomic chunks."""

    def test_load_docx_tables_as_atomic_chunks(self, tmp_path: Path):
        """Create a .docx with a heading, some text, and a table. Verify:
        - Table chunk has is_table: True in metadata
        - Table content is Markdown-formatted (with | and ---)
        - Table chunk's section matches the nearest preceding heading
        - Table chunk has correct section_level and section_path"""
        docx_path = tmp_path / "table_test.docx"
        doc = DocxDocument()

        # Add heading and text
        doc.add_heading("Employee Data", level=1)
        doc.add_paragraph("This section contains employee information.")

        # Add a table
        table = doc.add_table(rows=3, cols=2)
        table.rows[0].cells[0].text = "Name"
        table.rows[0].cells[1].text = "Department"
        table.rows[1].cells[0].text = "Alice"
        table.rows[1].cells[1].text = "Engineering"
        table.rows[2].cells[0].text = "Bob"
        table.rows[2].cells[1].text = "Marketing"

        doc.save(docx_path)

        result = load_docx(docx_path)

        # Find table chunks
        table_chunks = [d for d in result if d.metadata.get("is_table") is True]

        assert len(table_chunks) == 1, f"Expected 1 table chunk, got {len(table_chunks)}"

        table_chunk = table_chunks[0]

        # Verify is_table metadata
        assert table_chunk.metadata["is_table"] is True

        # Verify Markdown formatting
        content = table_chunk.page_content
        assert "|" in content, "Table should contain pipe characters"
        assert "---" in content, "Table should contain separator line"
        assert "Name" in content
        assert "Department" in content
        assert "Alice" in content
        assert "Engineering" in content

        # Verify section metadata matches preceding heading
        assert table_chunk.metadata["section"] == "Employee Data"
        assert table_chunk.metadata["section_level"] == 1
        assert table_chunk.metadata["section_path"] == "Employee Data"

    def test_load_docx_table_not_split(self, tmp_path: Path):
        """Create a .docx with a large table (many rows).
        Verify the table remains a single chunk (not split by the text splitter)."""
        docx_path = tmp_path / "large_table.docx"
        doc = DocxDocument()

        # Add heading
        doc.add_heading("Large Dataset", level=1)

        # Add a large table (50 rows)
        num_rows = 50
        table = doc.add_table(rows=num_rows, cols=3)

        # Header row
        table.rows[0].cells[0].text = "ID"
        table.rows[0].cells[1].text = "Name"
        table.rows[0].cells[2].text = "Description"

        # Data rows with substantial content
        for i in range(1, num_rows):
            table.rows[i].cells[0].text = str(i)
            table.rows[i].cells[1].text = f"Employee {i}"
            table.rows[i].cells[2].text = (
                f"This is a detailed description for employee number {i} with additional context"
            )

        doc.save(docx_path)

        result = load_docx(docx_path)

        # Find table chunks
        table_chunks = [d for d in result if d.metadata.get("is_table") is True]

        # Table should be a single chunk
        assert len(table_chunks) == 1, f"Expected 1 table chunk, got {len(table_chunks)}"

        # Verify the table content is substantial
        table_content = table_chunks[0].page_content
        assert "ID" in table_content
        assert "Name" in table_content
        assert "Description" in table_content
        # Should contain all rows
        assert "Employee 1" in table_content
        assert "Employee 49" in table_content or "Employee 48" in table_content


# ── load_xlsx column header tests ─────────────────────────────────


class TestLoadXlsxColumnHeaders:
    """Tests for column header extraction and preservation."""

    def test_load_xlsx_column_headers_preserved(self, tmp_path: Path):
        """Create an .xlsx with a header row and data rows. Verify:
        - Each chunk's page_content starts with the header row
        - column_headers metadata is present and correct"""
        xlsx_path = tmp_path / "headers_test.xlsx"

        sheet_data = {
            "Employees": [
                ("Name", "Department", "Salary"),
                ("Alice", "Engineering", "100000"),
                ("Bob", "Marketing", "80000"),
                ("Carol", "Sales", "90000"),
            ]
        }
        mock_wb = _create_mock_workbook(sheet_data)

        with patch("openpyxl.load_workbook") as mock_load:
            mock_load.return_value = mock_wb
            result = load_xlsx(xlsx_path)

        assert len(result) > 0

        for doc in result:
            # Verify column_headers metadata
            assert "column_headers" in doc.metadata
            assert doc.metadata["column_headers"] == "Name | Department | Salary"

            # Verify headers are prepended to content
            assert doc.page_content.startswith("Name | Department | Salary")

    def test_load_xlsx_empty_first_row_skipped(self, tmp_path: Path):
        """Create an .xlsx where the first row is empty.
        Verify the first non-empty row is used as headers."""
        xlsx_path = tmp_path / "empty_first_row.xlsx"

        sheet_data = {
            "Data": [
                (None, None, None),  # Empty first row
                ("", "", ""),  # Another empty row
                ("Product", "Price", "Quantity"),  # Actual header row
                ("Widget", "10.99", "100"),
                ("Gadget", "24.99", "50"),
            ]
        }
        mock_wb = _create_mock_workbook(sheet_data)

        with patch("openpyxl.load_workbook") as mock_load:
            mock_load.return_value = mock_wb
            result = load_xlsx(xlsx_path)

        assert len(result) > 0

        for doc in result:
            # Verify column_headers uses first non-empty row
            assert "column_headers" in doc.metadata
            assert doc.metadata["column_headers"] == "Product | Price | Quantity"

            # Verify content starts with the correct headers
            assert doc.page_content.startswith("Product | Price | Quantity")

