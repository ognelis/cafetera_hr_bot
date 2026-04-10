"""Tests for semantic chunking functionality in app.rag.parser."""

from __future__ import annotations

from pathlib import Path
from unittest.mock import MagicMock, patch

import pytest
from docx import Document as DocxDocument
from langchain_core.documents import Document as LCDocument
from langchain_core.embeddings import Embeddings

from app.config import Settings
from app.rag.parser import load_doc, load_document, load_docx


class FakeEmbeddings(Embeddings):
    """Mock embeddings class for testing semantic chunking.

    Returns deterministic embeddings based on hash of input text.
    """

    def embed_documents(self, texts: list[str]) -> list[list[float]]:
        return [[float(hash(t) % 100) / 100.0] * 10 for t in texts]

    def embed_query(self, text: str) -> list[float]:
        return [float(hash(text) % 100) / 100.0] * 10


# ── Fixtures ──────────────────────────────────────────────────────


@pytest.fixture
def mock_embeddings() -> FakeEmbeddings:
    """Provide a fake embeddings instance for testing."""
    return FakeEmbeddings()


@pytest.fixture
def multi_section_docx(tmp_path: Path) -> Path:
    """Create a .docx file with multiple topical sections for semantic chunking tests."""
    docx_path = tmp_path / "multi_section.docx"
    doc = DocxDocument()

    # Section 1: Introduction
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph(
        "This is the introduction section. It provides an overview of the document. "
        "The introduction sets the stage for what follows in later sections. "
        "Readers should understand the context and purpose after reading this part."
    )
    doc.add_paragraph(
        "The main goals are outlined here. Understanding these objectives is crucial "
        "for comprehending the rest of the material presented in this document."
    )

    # Section 2: Methodology
    doc.add_heading("Methodology", level=1)
    doc.add_paragraph(
        "The methodology section describes the approach taken. We use systematic "
        "analysis to evaluate different approaches. Data collection follows strict "
        "protocols to ensure reliability and validity of results."
    )
    doc.add_paragraph(
        "Statistical methods are applied to interpret findings. The analysis includes "
        "both quantitative and qualitative measures to provide comprehensive insights."
    )

    # Section 3: Results
    doc.add_heading("Results", level=1)
    doc.add_paragraph(
        "Results show significant improvements in efficiency. The findings demonstrate "
        "that the proposed approach outperforms existing methods. Performance metrics "
        "indicate a thirty percent reduction in processing time."
    )
    doc.add_paragraph(
        "Further analysis reveals additional benefits. These include improved accuracy "
        "and better resource utilization across all tested scenarios."
    )

    doc.save(docx_path)
    return docx_path


@pytest.fixture
def multi_paragraph_doc(tmp_path: Path) -> Path:
    """Provide a path to a .doc file with multi-paragraph content."""
    return tmp_path / "test_document.doc"


# ── Tests for load_document with recursive strategy (backward compatibility) ──


class TestLoadDocumentRecursiveDefault:
    """Tests confirming backward compatibility with recursive chunking."""

    def test_load_document_recursive_default(self, tmp_path: Path):
        """Calling load_document(path, strategy='recursive') on a .docx file produces
        chunks with correct source and section metadata."""
        docx_path = tmp_path / "test.docx"
        doc = DocxDocument()
        doc.add_heading("Test Section", level=1)
        doc.add_paragraph("This is test content for the recursive strategy.")
        doc.save(docx_path)

        result = load_document(docx_path, strategy="recursive")

        assert len(result) > 0
        assert all(isinstance(doc, LCDocument) for doc in result)
        for doc in result:
            assert doc.metadata["source"] == "test.docx"
            assert doc.metadata["section"] == "Test Section"


# ── Tests for load_document with semantic strategy ──


class TestLoadDocumentSemantic:
    """Tests for semantic chunking via load_document dispatcher."""

    def test_load_document_semantic_docx(
        self, multi_section_docx: Path, mock_embeddings: FakeEmbeddings
    ):
        """Calling load_document(path, strategy='semantic', embeddings=mock_embeddings)
        on a .docx file produces chunks with source and section metadata."""
        result = load_document(
            multi_section_docx,
            strategy="semantic",
            embeddings=mock_embeddings,
        )

        assert len(result) > 0
        assert all(isinstance(doc, LCDocument) for doc in result)
        assert all(len(doc.page_content) > 0 for doc in result)

        for doc in result:
            assert doc.metadata["source"] == "multi_section.docx"
            assert "section" in doc.metadata

    def test_load_document_semantic_doc(
        self, multi_paragraph_doc: Path, mock_embeddings: FakeEmbeddings
    ):
        """Calling load_document(path, strategy='semantic', embeddings=mock_embeddings)
        on a .doc file produces chunks with source and empty section metadata."""
        sample_text = (
            "First section about machine learning concepts and applications. "
            "Machine learning is a subset of artificial intelligence. "
            "It enables systems to learn and improve from experience.\n\n"
            "Second section covers deep learning methodologies. "
            "Deep learning uses neural networks with multiple layers. "
            "These networks can model complex patterns in data.\n\n"
            "Third section discusses natural language processing. "
            "NLP combines computational linguistics with machine learning. "
            "It enables computers to understand human language."
        )

        with patch("app.rag.parser.sharepoint2text.read_file") as mock_read_file:
            mock_result = MagicMock()
            mock_result.get_full_text.return_value = sample_text
            mock_read_file.return_value = iter([mock_result])
            result = load_document(
                multi_paragraph_doc,
                strategy="semantic",
                embeddings=mock_embeddings,
            )

        assert len(result) > 0
        assert all(isinstance(doc, LCDocument) for doc in result)
        assert all(len(doc.page_content) > 0 for doc in result)

        for doc in result:
            assert doc.metadata["source"] == "test_document.doc"
            assert doc.metadata["section"] == ""


# ── Tests for embeddings requirement validation ──


class TestSemanticStrategyRequiresEmbeddings:
    """Tests that semantic strategy requires embeddings parameter."""

    def test_load_document_semantic_requires_embeddings(self, tmp_path: Path):
        """Calling load_document(path, strategy='semantic') without embeddings
        raises ValueError."""
        docx_path = tmp_path / "test.docx"
        doc = DocxDocument()
        doc.add_heading("Test", level=1)
        doc.add_paragraph("Test content.")
        doc.save(docx_path)

        with pytest.raises(ValueError) as exc_info:
            load_document(docx_path, strategy="semantic")

        assert "embeddings are required" in str(exc_info.value)

    def test_load_docx_semantic_requires_embeddings(self, tmp_path: Path):
        """Calling load_docx(path, strategy='semantic') without embeddings
        raises ValueError."""
        docx_path = tmp_path / "test.docx"
        doc = DocxDocument()
        doc.add_heading("Test", level=1)
        doc.add_paragraph("Test content.")
        doc.save(docx_path)

        with pytest.raises(ValueError) as exc_info:
            load_docx(docx_path, strategy="semantic")

        assert "embeddings are required" in str(exc_info.value)

    def test_load_doc_semantic_requires_embeddings(self, multi_paragraph_doc: Path):
        """Calling load_doc(path, strategy='semantic') without embeddings
        raises ValueError."""
        with patch("app.rag.parser.sharepoint2text.read_file") as mock_read_file:
            mock_result = MagicMock()
            mock_result.get_full_text.return_value = "Some content here."
            mock_read_file.return_value = iter([mock_result])

            with pytest.raises(ValueError) as exc_info:
                load_doc(multi_paragraph_doc, strategy="semantic")

        assert "embeddings are required" in str(exc_info.value)


# ── Tests for Settings configuration defaults ──


class TestConfigChunkingDefaults:
    """Tests for chunking-related configuration defaults in Settings."""

    def test_config_chunk_strategy_default(self):
        """Settings defaults chunk_strategy to 'recursive'."""
        settings = Settings(_env_file=None)
        assert settings.chunk_strategy == "recursive"

    def test_config_semantic_defaults(self):
        """Settings defaults semantic_breakpoint_threshold_type to 'percentile'
        and semantic_breakpoint_threshold_amount to 95."""
        settings = Settings(_env_file=None)
        assert settings.semantic_breakpoint_threshold_type == "percentile"
        assert settings.semantic_breakpoint_threshold_amount == 95
