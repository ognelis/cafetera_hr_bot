"""Tests for Block 6 — RAG infrastructure."""

from __future__ import annotations

import sys
import types
from pathlib import Path
from tempfile import TemporaryDirectory
from unittest.mock import MagicMock, patch

from docx import Document as DocxDocument
from langchain_core.documents import Document as LCDocument

from app.config import Settings
from app.rag.chain import _format_docs, _format_docs_with_metadata, build_rag_chain
from app.rag.prompts import SYSTEM_PROMPT
from app.rag.retriever import COLLECTION_NAME, build_vectorstore, estimate_k

# ── Helpers ────────────────────────────────────────────────────────


def _make_docx(path: Path, paragraphs: list[tuple[str | None, str]]) -> None:
    """Create a .docx with paragraphs; if style is given, apply it."""
    doc = DocxDocument()
    for style, text in paragraphs:
        if style:
            doc.add_paragraph(text, style=style)
        else:
            doc.add_paragraph(text)
    doc.save(str(path))


# ── 6.0 — Config / Settings ───────────────────────────────────────


class TestRagSettings:
    def test_defaults_for_qdrant(self):
        s = Settings(vk_access_token="t", _env_file=None)
        assert s.qdrant_url == "http://localhost:6333"
        assert s.qdrant_api_key is None
        assert s.qdrant_collection == "hr_documents"

    def test_defaults_for_llm(self):
        s = Settings(vk_access_token="t", _env_file=None)
        assert s.llm_provider == "ollama"
        assert s.llm_model == "qwen3.5:4b-q4_K_M"
        assert s.llm_base_url == "http://localhost:11434"
        assert s.llm_api_key == ""

    def test_defaults_for_embeddings(self):
        s = Settings(vk_access_token="t", _env_file=None)
        assert s.embedding_model == "qwen3-embedding:4b-q4_K_M"

    def test_qdrant_settings_from_env(self, monkeypatch):
        monkeypatch.setenv("VK_ACCESS_TOKEN", "t")
        monkeypatch.setenv("QDRANT_URL", "http://qdrant:6333")
        monkeypatch.setenv("QDRANT_API_KEY", "secret")
        monkeypatch.setenv("QDRANT_COLLECTION", "custom_col")
        s = Settings()
        assert s.qdrant_url == "http://qdrant:6333"
        assert s.qdrant_api_key == "secret"
        assert s.qdrant_collection == "custom_col"

    def test_llm_settings_from_env(self, monkeypatch):
        monkeypatch.setenv("VK_ACCESS_TOKEN", "t")
        monkeypatch.setenv("LLM_PROVIDER", "openai")
        monkeypatch.setenv("LLM_MODEL", "gpt-4o-mini")
        monkeypatch.setenv("LLM_API_KEY", "sk-xxx")
        monkeypatch.setenv("LLM_BASE_URL", "https://api.openai.com/v1")
        s = Settings()
        assert s.llm_provider == "openai"
        assert s.llm_model == "gpt-4o-mini"
        assert s.llm_api_key == "sk-xxx"

    def test_llamacpp_provider_from_env(self, monkeypatch):
        monkeypatch.setenv("VK_ACCESS_TOKEN", "t")
        monkeypatch.setenv("LLM_PROVIDER", "llamacpp")
        monkeypatch.setenv("LLM_BASE_URL", "http://localhost:8080/v1")
        monkeypatch.setenv("LLM_API_KEY", "")
        s = Settings()
        assert s.llm_provider == "llamacpp"
        assert s.llm_base_url == "http://localhost:8080/v1"


# ── 6.1 — Ingestion: docx parsing ─────────────────────────────────


class TestExtractSections:
    def test_single_section_no_heading(self):
        from app.rag.parser import _extract_sections

        with TemporaryDirectory() as tmp:
            path = Path(tmp) / "test.docx"
            _make_docx(path, [(None, "Hello world"), (None, "Second line")])
            sections = _extract_sections(path)

        assert len(sections) == 1
        heading, body, level, section_path = sections[0]
        assert heading == ""
        assert "Hello world" in body
        assert "Second line" in body
        assert level == 0
        assert section_path == ""

    def test_multiple_headings(self):
        from app.rag.parser import _extract_sections

        with TemporaryDirectory() as tmp:
            path = Path(tmp) / "test.docx"
            _make_docx(
                path,
                [
                    ("Heading 1", "Chapter One"),
                    (None, "Body of chapter one."),
                    ("Heading 2", "Section Two"),
                    (None, "Body of section two."),
                ],
            )
            sections = _extract_sections(path)

        assert len(sections) == 2
        assert sections[0][0] == "Chapter One"
        assert "Body of chapter one." in sections[0][1]
        assert sections[0][2] == 1  # level
        assert sections[0][3] == "Chapter One"  # section_path
        assert sections[1][0] == "Section Two"
        assert "Body of section two." in sections[1][1]
        assert sections[1][2] == 2  # level
        assert sections[1][3] == "Chapter One > Section Two"  # section_path

    def test_empty_paragraphs_skipped(self):
        from app.rag.parser import _extract_sections

        with TemporaryDirectory() as tmp:
            path = Path(tmp) / "test.docx"
            _make_docx(path, [(None, "Text"), (None, "   "), (None, "More text")])
            sections = _extract_sections(path)

        assert len(sections) == 1
        assert "   " not in sections[0][1]
        assert sections[0][2] == 0  # level
        assert sections[0][3] == ""  # section_path


class TestLoadDocx:
    def test_returns_lc_documents(self):
        from app.rag.parser import load_docx

        with TemporaryDirectory() as tmp:
            path = Path(tmp) / "test.docx"
            _make_docx(path, [(None, "Some content for the HR bot.")])
            docs = load_docx(path)

        assert len(docs) >= 1
        assert all(isinstance(d, LCDocument) for d in docs)

    def test_metadata_contains_source(self):
        from app.rag.parser import load_docx

        with TemporaryDirectory() as tmp:
            path = Path(tmp) / "policy.docx"
            _make_docx(path, [(None, "Policy text here.")])
            docs = load_docx(path)

        assert docs[0].metadata["source"] == "policy.docx"

    def test_metadata_contains_section(self):
        from app.rag.parser import load_docx

        with TemporaryDirectory() as tmp:
            path = Path(tmp) / "manual.docx"
            _make_docx(
                path,
                [
                    ("Heading 1", "Vacation Policy"),
                    (None, "Employees are entitled to 28 days of leave."),
                ],
            )
            docs = load_docx(path)

        assert docs[0].metadata["section"] == "Vacation Policy"

    def test_long_text_is_chunked(self):
        from app.rag.parser import CHUNK_SIZE, load_docx

        # CHUNK_SIZE is in tokens (~500), so use ~2000 tokens to ensure splitting
        long_text = "Word " * 2000  # ~10000 chars, ~2000 tokens
        with TemporaryDirectory() as tmp:
            path = Path(tmp) / "long.docx"
            _make_docx(path, [(None, long_text)])
            docs = load_docx(path)

        assert len(docs) > 1
        for doc in docs:
            # CHUNK_SIZE is tokens; use a safe char upper bound (tokens rarely exceed 10 chars)
            assert len(doc.page_content) <= CHUNK_SIZE * 10


# ── 6.2 — Prompts ─────────────────────────────────────────────────


class TestSystemPrompt:
    def test_prompt_is_nonempty(self):
        assert len(SYSTEM_PROMPT) > 100

    def test_prompt_contains_context_placeholder(self):
        assert "{context}" in SYSTEM_PROMPT

    def test_prompt_mentions_hr(self):
        assert "HR" in SYSTEM_PROMPT

    def test_prompt_mentions_no_personal_data(self):
        lower = SYSTEM_PROMPT.lower()
        assert "персональные данные" in lower or "конфиденциальн" in lower

    def test_prompt_language_is_russian(self):
        assert "русском" in SYSTEM_PROMPT.lower()


# ── 6.2 — Chain helpers ───────────────────────────────────────────


class TestFormatDocs:
    def test_single_document(self):
        docs = [LCDocument(page_content="Hello")]
        assert _format_docs(docs) == "Hello"

    def test_multiple_documents_joined_with_separator(self):
        docs = [LCDocument(page_content="A"), LCDocument(page_content="B")]
        result = _format_docs(docs)
        assert "A" in result
        assert "B" in result
        assert "---" in result

    def test_empty_list(self):
        assert _format_docs([]) == ""


class TestFormatDocsWithMetadata:
    def test_single_document_with_filename(self):
        docs = [LCDocument(page_content="Hello", metadata={"filename": "test.docx"})]
        result = _format_docs_with_metadata(docs)
        assert "[Документ: test.docx]" in result
        assert "Hello" in result

    def test_single_document_without_filename(self):
        docs = [LCDocument(page_content="Hello")]
        result = _format_docs_with_metadata(docs)
        assert result == "Hello"
        assert "[Документ:" not in result

    def test_single_document_with_empty_filename(self):
        docs = [LCDocument(page_content="Hello", metadata={"filename": ""})]
        result = _format_docs_with_metadata(docs)
        assert result == "Hello"
        assert "[Документ:" not in result

    def test_multiple_documents_with_different_filenames(self):
        docs = [
            LCDocument(page_content="First chunk", metadata={"filename": "doc1.docx"}),
            LCDocument(page_content="Second chunk", metadata={"filename": "doc2.docx"}),
        ]
        result = _format_docs_with_metadata(docs)
        assert "[Документ: doc1.docx]" in result
        assert "[Документ: doc2.docx]" in result
        assert "First chunk" in result
        assert "Second chunk" in result
        assert "---" in result

    def test_multiple_documents_mixed_metadata(self):
        docs = [
            LCDocument(page_content="With filename", metadata={"filename": "test.docx"}),
            LCDocument(page_content="Without filename"),
            LCDocument(page_content="Empty filename", metadata={"filename": ""}),
        ]
        result = _format_docs_with_metadata(docs)
        assert "[Документ: test.docx]\nWith filename" in result
        assert "\nWithout filename" in result
        assert "\nEmpty filename" in result
        assert "---" in result

    def test_empty_list(self):
        assert _format_docs_with_metadata([]) == ""


# ── 6.2 — Retriever / vectorstore ─────────────────────────────────


class TestEstimateK:
    """Tests for adaptive k-tuning based on question complexity."""

    def test_short_questions_return_k2(self):
        """Questions with ≤5 words should return k=2."""
        assert estimate_k("Привет") == 2
        assert estimate_k("Как дела") == 2
        assert estimate_k("Что такое отпуск") == 2
        assert estimate_k("Как получить зарплату") == 2
        assert estimate_k("Сколько дней отпуска") == 2

    def test_medium_questions_return_k4(self):
        """Questions with 6-15 words should return k=4."""
        assert estimate_k("Как оформить отпуск в этой компании") == 4
        assert estimate_k("Что делать если я заболел и не могу работать") == 4
        assert estimate_k("Какие документы нужны для приема на работу сегодня") == 4
        assert estimate_k("Расскажи про политику компании по удаленной работе") == 4

    def test_long_questions_return_k6(self):
        """Questions with >15 words should return k=6."""
        # 16 words
        assert (
            estimate_k(
                "one two three four five six seven eight nine ten "
                "eleven twelve thirteen fourteen fifteen sixteen"
            )
            == 6
        )
        # 17 words
        assert (
            estimate_k(
                "Что мне делать если я хочу оформить отпуск "
                "но у меня уже есть планированный отпуск на следующий месяц"
            )
            == 6
        )

    def test_empty_string_returns_k2(self):
        """Empty string has 0 words, should return k=2."""
        assert estimate_k("") == 2

    def test_whitespace_only_returns_k2(self):
        """String with only whitespace has 0 words, should return k=2."""
        assert estimate_k("   ") == 2
        assert estimate_k("\t\n") == 2


class TestCollectionName:
    def test_default_collection_name(self):
        assert COLLECTION_NAME == "hr_documents"


class TestBuildVectorstore:
    def test_creates_vectorstore_instance(self):
        mock_client = MagicMock()
        mock_embeddings = MagicMock()

        with patch("app.rag.retriever.QdrantVectorStore") as mock_vs_cls:
            build_vectorstore(mock_client, mock_embeddings, "test_col")
            mock_vs_cls.assert_called_once_with(
                client=mock_client,
                collection_name="test_col",
                embedding=mock_embeddings,
            )


# ── 6.2 — Chain building ──────────────────────────────────────────


class TestBuildRagChain:
    def test_returns_runnable(self):
        mock_retriever = MagicMock()
        mock_llm = MagicMock()

        chain = build_rag_chain(mock_retriever, mock_llm, system_prompt=SYSTEM_PROMPT)
        # The chain should be a LangChain Runnable (has invoke method)
        assert hasattr(chain, "invoke")


# ── 6.3 — Llama.cpp provider dispatch ─────────────────────────────


class TestBuildLlmLlamaCpp:
    def _settings(self, **overrides):
        defaults = dict(
            vk_access_token="t",
            llm_provider="llamacpp",
            llm_model="test-model",
            llm_base_url="",
            llm_api_key="",
            _env_file=None,
        )
        defaults.update(overrides)
        return Settings(**defaults)

    def _fake_langchain_openai(self):
        mod = types.ModuleType("langchain_openai")
        mod.ChatOpenAI = MagicMock()
        mod.OpenAIEmbeddings = MagicMock()
        return mod

    def test_uses_chat_openai_with_defaults(self):
        from app.rag.chain import build_llm

        fake = self._fake_langchain_openai()
        s = self._settings()
        with patch.dict(sys.modules, {"langchain_openai": fake}):
            build_llm(s)
            fake.ChatOpenAI.assert_called_once_with(
                model="test-model",
                api_key="no-key",
                base_url="http://localhost:8080/v1",
                temperature=0.3,
            )

    def test_passes_real_api_key_when_set(self):
        from app.rag.chain import build_llm

        fake = self._fake_langchain_openai()
        s = self._settings(llm_api_key="real-key")
        with patch.dict(sys.modules, {"langchain_openai": fake}):
            build_llm(s)
            fake.ChatOpenAI.assert_called_once_with(
                model="test-model",
                api_key="real-key",
                base_url="http://localhost:8080/v1",
                temperature=0.3,
            )

    def test_passes_custom_base_url(self):
        from app.rag.chain import build_llm

        fake = self._fake_langchain_openai()
        s = self._settings(llm_base_url="http://gpu-box:9090/v1")
        with patch.dict(sys.modules, {"langchain_openai": fake}):
            build_llm(s)
            fake.ChatOpenAI.assert_called_once_with(
                model="test-model",
                api_key="no-key",
                base_url="http://gpu-box:9090/v1",
                temperature=0.3,
            )

    def test_import_error_message(self):
        import builtins

        from app.rag.chain import build_llm

        real_import = builtins.__import__

        def _block_openai(name, *args, **kwargs):
            if name == "langchain_openai":
                raise ImportError("no module")
            return real_import(name, *args, **kwargs)

        s = self._settings()
        with patch("builtins.__import__", side_effect=_block_openai):
            try:
                build_llm(s)
                raise AssertionError("Expected ImportError")  # noqa: TRY301
            except ImportError as exc:
                assert "openai_compatible" in str(exc)


class TestBuildEmbeddingsLlamaCpp:
    def _settings(self, **overrides):
        defaults = dict(
            vk_access_token="t",
            embedding_provider="llamacpp",
            embedding_model="test-embed",
            embedding_base_url="",
            embedding_api_key="",
            _env_file=None,
        )
        defaults.update(overrides)
        return Settings(**defaults)

    def _fake_langchain_openai(self):
        mod = types.ModuleType("langchain_openai")
        mod.ChatOpenAI = MagicMock()
        mod.OpenAIEmbeddings = MagicMock()
        return mod

    def test_uses_openai_embeddings_with_defaults(self):
        from app.rag.retriever import build_embeddings

        fake = self._fake_langchain_openai()
        s = self._settings()
        with patch.dict(sys.modules, {"langchain_openai": fake}):
            build_embeddings(s)
            fake.OpenAIEmbeddings.assert_called_once_with(
                model="test-embed",
                openai_api_key="no-key",
                openai_api_base="http://localhost:8080/v1",
            )

    def test_passes_real_api_key_when_set(self):
        from app.rag.retriever import build_embeddings

        fake = self._fake_langchain_openai()
        s = self._settings(embedding_api_key="real-key")
        with patch.dict(sys.modules, {"langchain_openai": fake}):
            build_embeddings(s)
            fake.OpenAIEmbeddings.assert_called_once_with(
                model="test-embed",
                openai_api_key="real-key",
                openai_api_base="http://localhost:8080/v1",
            )

    def test_import_error_message(self):
        import builtins

        from app.rag.retriever import build_embeddings

        real_import = builtins.__import__

        def _block_openai(name, *args, **kwargs):
            if name == "langchain_openai":
                raise ImportError("no module")
            return real_import(name, *args, **kwargs)

        s = self._settings()
        with patch("builtins.__import__", side_effect=_block_openai):
            try:
                build_embeddings(s)
                raise AssertionError("Expected ImportError")  # noqa: TRY301
            except ImportError as exc:
                assert "openai_compatible" in str(exc)
